"""
RackTrack — Setup Guide generator.

Produces docs/SETUP_GUIDE.docx — a clean, shareable Word document that walks a
new engineer through cloning, installing, and running the full RackTrack stack
on a fresh machine.

Re-run:
    python docs/_tools/generate_setup_guide.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Theme (same palette as docs/features/) ────────────────────────────────
NAVY        = RGBColor(0x1F, 0x3B, 0x6C)
NAVY_DARK   = RGBColor(0x10, 0x22, 0x44)
SLATE       = RGBColor(0x47, 0x55, 0x69)
INK         = RGBColor(0x1E, 0x29, 0x3B)
MUTED       = RGBColor(0x94, 0x9D, 0xAB)

HAIRLINE    = "CBD5E1"
HEADER_RULE = "94A3B8"
ACCENT_FILL = "1F3B6C"
CALLOUT_BG  = "F8FAFC"

BODY_FONT    = "Calibri"
HEADING_FONT = "Aptos Display"

OUT = Path(__file__).resolve().parent.parent / "SETUP_GUIDE.docx"


# ─── XML helpers ───────────────────────────────────────────────────────────
def _shade(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_border(cell, color=HAIRLINE, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), size); b.set(qn("w:color"), color)
        tcb.append(b)
    tc_pr.append(tcb)


def _clear_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}"); b.set(qn("w:val"), "nil")
        tcb.append(b)
    tc_pr.append(tcb)


def _row_h(row, twentieths: int) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(twentieths)); h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def _track(run, val: int = 30) -> None:
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:val"), str(val))
    rPr.append(spacing)


# ─── Page setup + base style ───────────────────────────────────────────────
def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    s.header_distance = Cm(1.0); s.footer_distance = Cm(1.0)

    n = doc.styles["Normal"]
    n.font.name = BODY_FONT
    n.font.size = Pt(10.5)
    n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.line_spacing = 1.32


# ─── Building blocks ───────────────────────────────────────────────────────
def title(doc: Document, eyebrow: str, title_text: str, tagline: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(eyebrow.upper())
    r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = SLATE
    _track(r, 40)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = NAVY_DARK
    r.font.name = HEADING_FONT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(tagline)
    r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(8)

    # title underline — slim slate hairline
    rule = doc.add_table(rows=1, cols=1)
    cell = rule.rows[0].cells[0]; _shade(cell, HEADER_RULE); _clear_border(cell)
    _row_h(rule.rows[0], 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = NAVY
    _track(r, 40)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    rule = doc.add_table(rows=1, cols=1)
    cell = rule.rows[0].cells[0]; _shade(cell, HAIRLINE); _clear_border(cell)
    _row_h(rule.rows[0], 6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph().add_run(text)


def callout(doc: Document, label: str, text: str) -> None:
    """Same single-cell, navy-left-rule callout as the Feature Reference docs."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, CALLOUT_BG)
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:color"), ACCENT_FILL)
    tcb.append(left)
    for edge in ("top", "right", "bottom"):
        b = OxmlElement(f"w:{edge}"); b.set(qn("w:val"), "nil")
        tcb.append(b)
    tc_pr.append(tcb)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(label + "   ")
    r.font.bold = True; r.font.color.rgb = NAVY_DARK; r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.color.rgb = INK; r2.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def bullets(doc: Document, items: list) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            head, rest = it
            r = p.add_run(head); r.font.bold = True; r.font.color.rgb = NAVY_DARK
            p.add_run(" — " + rest)
        else:
            p.add_run(it)


def numbered(doc: Document, items: list) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(it)


def table(doc: Document, headers: list[str], rows: list[list[str]],
          widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _set_border(cell, color=HAIRLINE)
        _shade(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.15)
        r = p.add_run(h.upper())
        r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(9)
        _track(r, 30)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            _set_border(cell, color=HAIRLINE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.15)
            r = p.add_run(val); r.font.size = Pt(10); r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def code(doc: Document, snippet: str) -> None:
    """Monospace code block, soft-fill background."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _set_border(cell, color=HAIRLINE)
    _shade(cell, "F1F5F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.2)
    r = p.add_run(snippet)
    r.font.name = "Consolas"; r.font.size = Pt(10); r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def footer(doc: Document) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    s = doc.sections[0]

    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(17.2), WD_TAB_ALIGNMENT.RIGHT)
    rh = hp.add_run("RackTrack")
    rh.font.size = Pt(8); rh.font.bold = True; rh.font.color.rgb = NAVY
    _track(rh, 40)
    rh2 = hp.add_run("   ·   Setup Guide")
    rh2.font.size = Pt(8); rh2.font.color.rgb = SLATE
    hp.add_run("\t")
    rh3 = hp.add_run("Onboarding")
    rh3.font.size = Pt(8); rh3.font.italic = True; rh3.font.color.rgb = MUTED

    rule = s.header.add_table(rows=1, cols=1, width=Cm(17.2))
    rc = rule.rows[0].cells[0]; _shade(rc, HAIRLINE); _clear_border(rc)
    _row_h(rule.rows[0], 8)

    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(8.6), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.2), WD_TAB_ALIGNMENT.RIGHT)
    f1 = fp.add_run("RackTrack — Setup Guide")
    f1.font.size = Pt(8); f1.font.color.rgb = MUTED
    fp.add_run("\t")
    f2 = fp.add_run("For new engineers")
    f2.font.size = Pt(8); f2.font.color.rgb = MUTED
    fp.add_run("\t")
    f3 = fp.add_run(); f3.font.size = Pt(8); f3.font.color.rgb = MUTED
    fld_b = OxmlElement("w:fldChar"); fld_b.set(qn("w:fldCharType"), "begin")
    fld_i = OxmlElement("w:instrText"); fld_i.text = "PAGE"
    fld_e = OxmlElement("w:fldChar"); fld_e.set(qn("w:fldCharType"), "end")
    f3._r.append(fld_b); f3._r.append(fld_i); f3._r.append(fld_e)


# ─── Build the document ────────────────────────────────────────────────────
def build() -> None:
    doc = Document()
    setup(doc)

    title(doc,
        "RackTrack  ·  Onboarding",
        "Setup Guide",
        "Clone the repo and bring up the full stack on a fresh machine.")

    callout(doc, "Read me first.",
        "This guide walks a new engineer through getting RackTrack running locally — "
        "from installing prerequisites to capturing a first scan. Allow about 30–45 "
        "minutes end-to-end on a fast connection (the Git LFS pull alone is ~2.2 GB).")

    # ── 1. Prerequisites ──
    section(doc, "1. Prerequisites (install once)")
    para(doc, "Everything below must be available on the PATH before you clone the repo. The exact installation method depends on your OS; the commands shown assume Windows with winget.")
    table(doc,
          headers=["Tool", "Why it's needed", "Install (Windows)"],
          rows=[
              ["Node.js 18 or newer", "Client + server runtime", "winget install OpenJS.NodeJS.LTS"],
              ["Python 3.10 or newer", "Pipeline, agents, ServiceNow bridge", "winget install Python.Python.3.12"],
              ["Git LFS", "Required — model weights (~2.2 GB) are stored in LFS", "winget install -e --id GitHub.GitLFS"],
              ["Android Studio (optional)", "Only needed for Android app builds", "winget install Google.AndroidStudio"],
              ["Xcode (optional, Mac only)", "Only needed for iOS app builds", "Mac App Store"],
              ["Docker Desktop (optional)", "Only needed for the bundled Netdisco stack", "winget install Docker.DockerDesktop"],
          ])
    para(doc, "After installing Git LFS, run this one-time initialisation in any PowerShell window:")
    code(doc, "git lfs install")

    # ── 2. Clone ──
    section(doc, "2. Clone the repository")
    para(doc, "Clone the production repo. LFS files (model weights) download automatically during the clone, provided Git LFS is installed.")
    code(doc,
         "git clone https://github.com/aasrithasravani016-ui/RACKTRACK_FINAL_V1.git\n"
         "cd RACKTRACK_FINAL_V1")
    para(doc, "If model files appear as tiny text \"pointer\" files (typically a few hundred bytes) instead of the real binaries, pull them explicitly:")
    code(doc, "git lfs pull")

    # ── 3. Install dependencies ──
    section(doc, "3. Install dependencies")
    para(doc, "One Node command installs everything across root, client, and server. Then run the Python pip commands for the agents and pipeline.")
    code(doc, "npm run install:all")
    code(doc,
         "pip install -r requirements.txt\n"
         "pip install -r servicenow/requirements.txt\n"
         "pip install -r Agent/Agent_scrap/requirements.txt\n"
         "pip install -r s_agent/requirements.txt")
    para(doc, "It is good practice to install the Python dependencies into a virtual environment to avoid polluting the system Python, though the project does not strictly require it.")

    # ── 4. Environment files ──
    section(doc, "4. Set up environment files")
    para(doc, "Most environment files are committed with structure but with sensitive values removed. You will need to fill in tenant-specific secrets before features that depend on them will work.")

    bullets(doc, [
        ("server/.env",
         "Server configuration, JWT signing secret, switch SSH credentials. "
         "Regenerate server/data/jwt.secret on every new install — never reuse "
         "JWT signing keys across environments."),
        ("servicenow/.env",
         "ServiceNow instance URL and credentials. Use servicenow/.env.example as the template; fill in your PDI or production instance details."),
        ("client/.env.production",
         "Client-side configuration that ships in production builds."),
        ("Atlassian API token (only if publishing docs to Confluence)",
         "Create a token at https://id.atlassian.com/manage-profile/security/api-tokens "
         "and add ATLASSIAN_API_TOKEN=ATATT-… to server/.env on machines that need it. "
         "Do not commit the token."),
    ])

    callout(doc, "Security note.",
        "Treat every .env file as sensitive. Never commit secrets, never paste tokens into "
        "chat or shared documents, and rotate tokens promptly if you suspect they have been "
        "exposed.")

    # ── 5. Run ──
    section(doc, "5. Start the stack")
    para(doc, "Two options. Use dev mode for day-to-day local development; use the production-style launcher when you want to test on a phone over the Cloudflare quick-tunnel.")

    para(doc, "Dev mode (client + server with hot reload):")
    code(doc, "npm run dev")

    para(doc, "Production-style on Windows (server plus a Cloudflare quick-tunnel for phone access):")
    code(doc, ".\\start.ps1")
    para(doc, "When the tunnel is up, its public URL is written to current-url.txt at the repo root. The helper script update-apk-url.ps1 patches the Android build to point at that URL.")

    # ── 6. Mobile builds ──
    section(doc, "6. Mobile builds (only if you are shipping to a device)")
    para(doc, "The client is a React + Capacitor app. To produce native iOS or Android builds, build the web assets and then sync them into the native projects.")
    code(doc,
         "cd client\n"
         "npm run build\n"
         "npx cap sync android         # or: npx cap sync ios\n"
         "npx cap open android         # opens Android Studio")
    para(doc, "On Android, the camera and AR sub-mode require ARCore-capable hardware (Pixel, recent Galaxy S, etc.); lower-end devices fall back to the standard capture flow.")

    # ── 7. Netdisco ──
    section(doc, "7. Optional — the Netdisco integration")
    para(doc, "RackTrack ships a bundled Netdisco Docker stack for demos and local development. Spin it up only if you need the live-network integration available.")
    code(doc,
         "cd netdisco-docker\n"
         "docker compose up -d")
    para(doc, "The bundled stack is for development use only; production environments should connect to a long-lived, separately-managed Netdisco instance.")

    # ── 8. Sanity check ──
    section(doc, "8. Sanity check the install")
    para(doc, "After everything is installed and running, verify the basics:")
    code(doc,
         "node --version              # >= 18\n"
         "python --version            # >= 3.10\n"
         "git lfs ls-files | head     # should list .pt / .onnx model files\n"
         "npm run dev                 # client should open on http://localhost:5173")
    para(doc, "If the server logs \"listening on port 3001\" and the client opens its home screen, the install is correct. Capture a single rack photo to confirm the full pipeline works end-to-end.")

    # ── 9. What's deliberately not in the repo ──
    section(doc, "9. What is deliberately not in the repository")
    table(doc,
          headers=["Item", "Why not in repo", "How to handle"],
          rows=[
              ["node_modules/", "Build artifact, gitignored", "Recreated automatically by npm run install:all"],
              ["outputs/", "Runtime scan records, gitignored", "Created automatically as scans run"],
              ["server/uploads/", "Runtime upload storage, gitignored", "Created automatically"],
              ["client/public.zip", "Build artifact, excluded from production commit", "Re-derive from client/public/ if needed"],
              ["Atlassian API token", "Sensitive — never committed", "Recreate per machine that needs to publish docs"],
              ["JWT signing secret", "Sensitive — regenerated per environment", "Regenerate server/data/jwt.secret on every new install"],
          ])

    # ── 10. Common problems ──
    section(doc, "10. If something goes wrong")
    bullets(doc, [
        ("Model files are tiny",
         "You likely cloned without Git LFS installed. Run \"git lfs install\" and then \"git lfs pull\" to fetch the real model binaries."),
        ("npm install fails on a native dependency",
         "Make sure you have Visual Studio Build Tools (Windows) or Xcode Command Line Tools (Mac) installed; some Capacitor native modules need a C++ toolchain."),
        ("Server starts but client can't reach it",
         "Default ports are 3001 (server) and 5173 (Vite dev). Check both are free and not blocked by a firewall."),
        ("ServiceNow features show as unavailable",
         "Confirm servicenow/.env points at a reachable instance and the credentials have read access to the incident and CMDB tables."),
        ("AR sub-mode missing on Android",
         "ARCore support is device-dependent. Lower-end devices automatically hide the AR tab and fall back to Photo / Video capture."),
        ("PDF report generation fails",
         "PDF rendering needs a headless Chromium. Re-install Puppeteer dependencies inside the server folder."),
    ])

    # ── 11. Where to go next ──
    section(doc, "11. Where to go next")
    para(doc, "Once the stack is running locally, the recommended reading order is:")
    numbered(doc, [
        "docs/SETUP_GUIDE.docx — this document (keep handy).",
        "RackTrack Confluence space → \"Product Overview\" container, for the executive view of what the platform does.",
        "RackTrack Confluence space → \"Feature Reference\" container, for the engineering-level detail on every backend feature.",
        "docs/features/ in the repo — the same Feature Reference content as Word documents, in case you prefer Word over Confluence.",
    ])

    footer(doc)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
